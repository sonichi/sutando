#!/usr/bin/env python3
"""doc-ingest: extract readable text from document files (PDF/XLSX/CSV/DOCX/PPTX/...).

Best-available extraction with graceful fallbacks — works with stdlib alone,
uses poppler/openpyxl/python-docx/textutil when present. See SKILL.md.
"""
from __future__ import annotations

import csv
import io
import json
import os
import re
import selectors
import shutil
import subprocess
import sys
import time
import zipfile
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Optional

DEFAULT_MAX_CHARS = 200_000
DEFAULT_MAX_ROWS = 500
MAX_ARCHIVE_DEPTH = 8
MAX_TABLE_BYTES = 64 * 1024 * 1024
MAX_TABLE_ROWS = 100_000
MAX_TABLE_CELLS = 1_000_000
MAX_TABLE_TEXT_CHARS = 16 * 1024 * 1024
# --csv compute-exact view: source/output ceilings (fail-closed, never truncate).
CSV_MAX_SOURCE_BYTES = 32 * 1024 * 1024
CSV_MAX_OUTPUT_BYTES = 64 * 1024 * 1024
MAX_XLSX_COLUMNS = 16_384
TEXT_READ_CHUNK_CHARS = 64 * 1024
MAX_DOCUMENT_TEXT_CHARS = 16 * 1024 * 1024
MAX_DOCUMENT_BYTES = 64 * 1024 * 1024
TEXT_COMMAND_TIMEOUT_SECONDS = 120

IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".heic"}
AUDIO_SUFFIXES = {".mp3", ".wav", ".m4a", ".ogg", ".flac", ".aac", ".opus"}
TEXT_SUFFIXES = {
    ".txt", ".md", ".markdown", ".json", ".jsonl", ".xml", ".html", ".htm",
    ".yaml", ".yml", ".toml", ".ini", ".log", ".py", ".js", ".ts", ".sh",
    ".c", ".cpp", ".h", ".java", ".go", ".rs", ".rb", ".sql",
}


def _truncate(text: str, max_chars: int) -> str:
    if isinstance(text, _PreTruncatedText):
        return text
    if max_chars and len(text) > max_chars:
        return text[:max_chars] + f"\n\n[doc-ingest: truncated at {max_chars} chars — original {len(text)} chars]"
    return text


class _PreTruncatedText(str):
    """A bounded stream result that already carries the exact truncation notice."""


class _ResourceBudgetError(RuntimeError):
    """Extraction stopped before untrusted input could exceed a hard bound."""


def _document_budgets(max_chars: int) -> tuple[int, int]:
    """Return hard character/byte ceilings for untrusted document extraction."""
    char_budget = MAX_DOCUMENT_TEXT_CHARS if not max_chars else min(
        MAX_DOCUMENT_TEXT_CHARS, max(max_chars * 4, TEXT_READ_CHUNK_CHARS)
    )
    return char_budget, min(MAX_DOCUMENT_BYTES, max(char_budget * 4, 1024 * 1024))


def _join_bounded(chunks, max_chars: int, label: str, separator: str = "\n") -> str:
    """Join incrementally while refusing materialization past the hard ceiling."""
    char_budget, _ = _document_budgets(max_chars)
    parts: list[str] = []
    total = 0
    for chunk in chunks:
        chunk = str(chunk)
        total += len(chunk)
        if total > char_budget:
            raise _ResourceBudgetError(
                f"{label} exceeds extracted text budget ({char_budget} chars); extraction stopped"
            )
        parts.append(chunk)
    return separator.join(parts)


def _run_text_command_bounded(command: list[str], max_chars: int, label: str) -> tuple[int, str, str]:
    """Stream converter output and stop the child as soon as its budget is crossed."""
    _, byte_budget = _document_budgets(max_chars)
    proc = subprocess.Popen(command, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    assert proc.stdout is not None and proc.stderr is not None
    streams = selectors.DefaultSelector()
    streams.register(proc.stdout, selectors.EVENT_READ, "stdout")
    streams.register(proc.stderr, selectors.EVENT_READ, "stderr")
    stdout_chunks: list[bytes] = []
    stderr_chunks: list[bytes] = []
    stdout_size = 0
    stderr_size = 0
    deadline = time.monotonic() + TEXT_COMMAND_TIMEOUT_SECONDS

    def stop_child() -> None:
        if proc.poll() is None:
            proc.terminate()
            try:
                proc.wait(timeout=2)
            except subprocess.TimeoutExpired:
                proc.kill()
                proc.wait()

    try:
        while streams.get_map():
            remaining = deadline - time.monotonic()
            if remaining <= 0:
                stop_child()
                raise subprocess.TimeoutExpired(command, TEXT_COMMAND_TIMEOUT_SECONDS)
            for key, _ in streams.select(timeout=min(remaining, 0.1)):
                # Read at most one byte beyond the stdout budget. This makes the
                # crossing observable without draining a child that keeps
                # producing output; the child is terminated immediately below.
                if key.data == "stdout":
                    read_size = min(64 * 1024, byte_budget - stdout_size + 1)
                else:
                    read_size = 64 * 1024
                chunk = os.read(key.fd, read_size)
                if not chunk:
                    streams.unregister(key.fileobj)
                    continue
                if key.data == "stdout":
                    stdout_size += len(chunk)
                    if stdout_size > byte_budget:
                        stop_child()
                        raise _ResourceBudgetError(
                            f"{label} exceeds extracted byte budget "
                            f"({byte_budget} bytes); extraction stopped"
                        )
                    stdout_chunks.append(chunk)
                elif stderr_size < 64 * 1024:
                    keep = chunk[:64 * 1024 - stderr_size]
                    stderr_chunks.append(keep)
                    stderr_size += len(keep)
        returncode = proc.wait(timeout=max(0.1, deadline - time.monotonic()))
    finally:
        streams.close()
        stop_child()
        proc.stdout.close()
        proc.stderr.close()
    return (
        returncode,
        b"".join(stdout_chunks).decode("utf-8", errors="replace"),
        b"".join(stderr_chunks).decode("utf-8", errors="replace"),
    )


def _read_ooxml_member_bounded(zf: zipfile.ZipFile, name: str, max_chars: int) -> bytes:
    """Reject oversized OOXML members before zipfile allocates their payload."""
    _, byte_budget = _document_budgets(max_chars)
    info = zf.getinfo(name)
    if info.file_size > byte_budget:
        raise _ResourceBudgetError(
            f"OOXML member exceeds uncompressed byte budget ({byte_budget} bytes); extraction stopped"
        )
    return zf.read(info)


def _validate_ooxml_archive_bounded(zf: zipfile.ZipFile, max_chars: int) -> None:
    """Reject archive inflation before a third-party OOXML parser opens it."""
    _, byte_budget = _document_budgets(max_chars)
    total = 0
    for info in zf.infolist():
        total += info.file_size
        if total > byte_budget:
            raise _ResourceBudgetError(
                f"OOXML archive exceeds uncompressed byte budget ({byte_budget} bytes); extraction stopped"
            )


def _read_text_bounded(path: Path, max_chars: int) -> str:
    """Decode text incrementally, retaining at most the requested output prefix."""
    if not max_chars:
        return path.read_text(encoding="utf-8", errors="replace")
    kept: list[str] = []
    kept_chars = 0
    total_chars = 0
    with path.open("r", encoding="utf-8", errors="replace") as fh:
        while True:
            chunk = fh.read(TEXT_READ_CHUNK_CHARS)
            if not chunk:
                break
            total_chars += len(chunk)
            if kept_chars < max_chars:
                prefix = chunk[:max_chars - kept_chars]
                kept.append(prefix)
                kept_chars += len(prefix)
    text = "".join(kept)
    if total_chars > max_chars:
        return _PreTruncatedText(
            text
            + f"\n\n[doc-ingest: truncated at {max_chars} chars — "
            f"original {total_chars} chars]"
        )
    return text


# A numeric cell whose exact decimal expansion would be enormous is not
# meaningful tabular data. Because the summary is built BEFORE _truncate() runs,
# rendering (or aggregating) such a value on an untrusted attachment is a
# resource-exhaustion vector — e.g. a cell of "1e-100000" expands to ~300k chars.
# Cap both the order of magnitude and the significant-digit count so any accepted
# number renders in O(cap) chars; anything past the cap falls back to plain text.
_NUM_MAGNITUDE_CAP = 1000


def _parse_number(cell: str):
    """Exact, *bounded* numeric parse of a cell → a finite Decimal, else None.

    Uses Decimal (not float) so integers beyond IEEE-754 precision (>2**53) keep
    their exact value instead of silently rounding. Rejects (→ None, treated as
    text) anything that isn't a safe numeric datum: Decimal ALSO parses 'NaN' /
    'Infinity' (→ reject via is_finite), and a finite value whose magnitude or
    digit count exceeds _NUM_MAGNITUDE_CAP is rejected so a degenerate exponent
    like '1e-100000' can't blow the summary up before truncation."""
    try:
        d = Decimal(cell.replace(",", ""))
    except InvalidOperation:
        return None
    if not d.is_finite():
        return None
    if abs(d.adjusted()) > _NUM_MAGNITUDE_CAP or len(d.as_tuple().digits) > _NUM_MAGNITUDE_CAP:
        return None
    return d


def _fmt_num(x: Decimal) -> str:
    """Exact decimal string for a finite Decimal; drop a trailing '.0' when integral
    and never emit exponent notation (fixed-point, so a solver reads a plain number)."""
    integral = x.to_integral_value()
    return str(integral) if x == integral else format(x.normalize(), "f")


class _TableCollector:
    """Incrementally render/digest an untrusted table within hard resource caps."""

    def __init__(self, max_rows: int, *, budget: Optional[dict[str, int]] = None):
        self.max_rows = max_rows
        self.budget = budget if budget is not None else {
            "rows": MAX_TABLE_ROWS,
            "cells": MAX_TABLE_CELLS,
            "text": MAX_TABLE_TEXT_CHARS,
        }
        self.total_rows = 0
        self.width = 0
        self.header: list[str] = []
        self.shown: list[list[str]] = []
        self.stats: list[dict] = []

    def add(self, raw_row) -> None:
        row = ["" if cell is None else str(cell) for cell in raw_row]
        row_text = sum(len(cell) for cell in row)
        if self.budget["rows"] <= 0:
            raise RuntimeError(
                f"table exceeds shared row budget ({MAX_TABLE_ROWS}); extraction stopped"
            )
        if len(row) > self.budget["cells"]:
            raise RuntimeError(
                f"table exceeds shared cell budget ({MAX_TABLE_CELLS}); extraction stopped"
            )
        if row_text > self.budget["text"]:
            raise RuntimeError(
                f"table exceeds shared text budget ({MAX_TABLE_TEXT_CHARS} chars); extraction stopped"
            )
        self.budget["rows"] -= 1
        self.budget["cells"] -= len(row)
        self.budget["text"] -= row_text
        self.total_rows += 1
        self.width = max(self.width, len(row))
        if not self.max_rows or len(self.shown) < self.max_rows:
            self.shown.append(row)
        if self.total_rows == 1:
            self.header = row
            return
        while len(self.stats) < len(row):
            self.stats.append({
                "count": 0,
                "numeric": True,
                "sum": Decimal(0),
                "min": None,
                "max": None,
            })
        for j, cell in enumerate(row):
            value = cell.strip()
            if not value:
                continue
            stat = self.stats[j]
            stat["count"] += 1
            if not stat["numeric"]:
                continue
            number = _parse_number(value)
            if number is None:
                stat["numeric"] = False
                continue
            stat["sum"] += number
            stat["min"] = number if stat["min"] is None else min(stat["min"], number)
            stat["max"] = number if stat["max"] is None else max(stat["max"], number)

    def summary(self) -> str:
        if self.total_rows < 2:
            return ""
        out = [
            f"**Table summary:** {self.total_rows - 1} data rows × {self.width} columns."
        ]
        for j in range(self.width):
            stat = self.stats[j] if j < len(self.stats) else {
                "count": 0, "numeric": False, "sum": Decimal(0), "min": None, "max": None,
            }
            name = (
                self.header[j].strip()
                if j < len(self.header) and self.header[j].strip()
                else f"col{j + 1}"
            )
            seg = f"- **{name}**: {stat['count']} non-empty"
            if stat["count"] and stat["numeric"]:
                seg += (
                    f"; numeric → sum {_fmt_num(stat['sum'])}, "
                    f"min {_fmt_num(stat['min'])}, max {_fmt_num(stat['max'])}"
                )
            out.append(seg)
        return "\n".join(out)

    def render(self, *, summary: bool = False) -> str:
        if not self.shown:
            return "(empty table)"
        shown = self.shown
        width = max(len(r) for r in shown)
        norm = [r + [""] * (width - len(r)) for r in shown]
        lines = ["| " + " | ".join(str(c) for c in norm[0]) + " |",
                 "|" + "---|" * width]
        lines += ["| " + " | ".join(str(c) for c in r) + " |" for r in norm[1:]]
        if self.total_rows > len(shown):
            lines.append(
                f"\n[doc-ingest: showing {len(shown)} of {self.total_rows} rows]"
            )
        table = "\n".join(lines)
        digest = self.summary() if summary else ""
        return digest + "\n\n" + table if digest else table


def _collect_table(rows, max_rows: int, *, summary: bool = False,
                   budget: Optional[dict[str, int]] = None) -> str:
    collector = _TableCollector(max_rows, budget=budget)
    for row in rows:
        collector.add(row)
    return collector.render(summary=summary)


def _table_summary(rows: list[list[str]]) -> str:
    """Computed structural digest over the full row stream."""
    collector = _TableCollector(0)
    for row in rows:
        collector.add(row)
    return collector.summary()


def _rows_to_markdown(rows: list[list[str]], max_rows: int, *, summary: bool = False) -> str:
    return _collect_table(rows, max_rows, summary=summary)


def _xml_text(payload: bytes) -> str:
    # Dependency-free OOXML text scrape: drop tags, keep text runs.
    text = payload.decode("utf-8", errors="replace")
    text = re.sub(r"<w:p[ >]", "\n<", text)  # paragraph boundaries → newlines
    text = re.sub(r"<a:p>", "\n<", text)
    text = re.sub(r"<[^>]+>", "", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def extract_pdf(path: Path, max_chars: int = DEFAULT_MAX_CHARS) -> str:
    tried = []
    if shutil.which("pdftotext"):
        tried.append("pdftotext")
        returncode, stdout, _ = _run_text_command_bounded(
            ["pdftotext", "-layout", str(path), "-"], max_chars, "PDF"
        )
        if returncode == 0 and stdout.strip():
            return stdout
    for mod, call in (
        ("pypdf", lambda m: _join_bounded(
            ((p.extract_text() or "") for p in m.PdfReader(str(path)).pages), max_chars, "PDF"
        )),
        ("fitz", lambda m: _join_bounded(
            (p.get_text() for p in m.open(str(path))), max_chars, "PDF"
        )),
    ):
        try:
            extractor = __import__(mod)
        except ImportError:
            continue
        tried.append(mod)
        try:
            text = call(extractor)
        except _ResourceBudgetError:
            raise
        except Exception:  # noqa: BLE001 — fall through to the next extractor
            continue
        if text.strip():
            return text
        # An extractor that returns no text has NOT succeeded — a no-text-layer
        # or image-only PDF must fall through to the next extractor and then
        # fail honestly, never return an empty string as a successful result.
    if tried:
        raise RuntimeError(f"PDF extraction failed (tried: {', '.join(tried)}) — no text extracted; file may be image-only or corrupt")
    raise RuntimeError("no PDF extractor available (need poppler's pdftotext, pypdf, or PyMuPDF)")


def _col_index(cell_ref: str, max_columns: int = MAX_XLSX_COLUMNS) -> int:
    """'B2' -> 1 (0-based column). Falls back to 0 if the ref has no letters."""
    letters = re.match(r"[A-Za-z]+", cell_ref or "")
    if not letters:
        return 0
    idx = 0
    for ch in letters.group(0).upper():
        idx = idx * 26 + (ord(ch) - 64)
        if idx > max_columns:
            raise RuntimeError(
                f"XLSX cell reference exceeds column/cell budget "
                f"({max_columns}); extraction stopped"
            )
    return idx - 1


def _xlsx_zip_fallback(path: Path, max_rows: int) -> str:
    """Dependency-free worksheet parse (no openpyxl): reads shared strings,
    inline strings, and numeric/raw cell values, placing each cell by its
    column reference so sparse rows stay aligned. This is the extraction the
    skill promises — the old shared-strings-only path silently dropped inline
    strings and numeric cells."""
    import xml.etree.ElementTree as ET  # noqa: PLC0415 — stdlib, lazy for import cost

    def local(tag: str) -> str:
        return tag.rsplit("}", 1)[-1]

    with zipfile.ZipFile(path) as zf:
        names = zf.namelist()
        table_names = [
            name for name in names
            if name == "xl/sharedStrings.xml"
            or re.fullmatch(r"xl/worksheets/sheet\d+\.xml", name)
        ]
        table_bytes = sum(zf.getinfo(name).file_size for name in table_names)
        if table_bytes > MAX_TABLE_BYTES:
            raise RuntimeError(
                f"XLSX table XML exceeds uncompressed byte budget "
                f"({MAX_TABLE_BYTES // (1024 * 1024)} MiB); extraction stopped"
            )
        shared: list[str] = []
        if "xl/sharedStrings.xml" in names:
            root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
            for si in root:
                shared.append("".join(t.text or "" for t in si.iter() if local(t.tag) == "t"))

        sheets = sorted(n for n in names if re.fullmatch(r"xl/worksheets/sheet\d+\.xml", n))
        parts: list[str] = []
        budget = {
            "rows": MAX_TABLE_ROWS,
            "cells": MAX_TABLE_CELLS,
            "text": MAX_TABLE_TEXT_CHARS,
        }
        for i, sheet_name in enumerate(sheets, 1):
            collector = _TableCollector(max_rows, budget=budget)
            with zf.open(sheet_name) as sheet_fh:
                row_iter = (
                    elem for _event, elem in ET.iterparse(sheet_fh, events=("end",))
                    if local(elem.tag) == "row"
                )
                for row in row_iter:
                    cells: dict[int, str] = {}
                    for c in row:
                        if local(c.tag) != "c":
                            continue
                        if budget["cells"] <= 0:
                            raise RuntimeError(
                                f"table exceeds shared cell budget "
                                f"({MAX_TABLE_CELLS}); extraction stopped"
                            )
                        col = _col_index(
                            c.get("r", ""),
                            min(MAX_XLSX_COLUMNS, budget["cells"]),
                        )
                        ctype = c.get("t")
                        if ctype == "s":  # shared string: <v> holds the index
                            v = next((x for x in c if local(x.tag) == "v"), None)
                            try:
                                val = shared[int(v.text)] if v is not None and v.text else ""
                            except (ValueError, IndexError):
                                val = ""
                        elif ctype == "inlineStr":  # inline: <is><t>…</t></is>
                            val = "".join(x.text or "" for x in c.iter() if local(x.tag) == "t")
                        else:  # numeric/boolean/raw: <v>
                            v = next((x for x in c if local(x.tag) == "v"), None)
                            val = (v.text or "") if v is not None else ""
                        cells[col] = val
                    width = max(cells) + 1 if cells else 0
                    collector.add([cells.get(j, "") for j in range(width)])
                    row.clear()
            body = collector.render(summary=True)
            parts.append(f"## Sheet {i}\n\n{body}")
        return "\n\n".join(parts) if parts else "(xlsx: no worksheet data found)"


def extract_xlsx(path: Path, max_rows: int) -> str:
    if path.stat().st_size > MAX_TABLE_BYTES:
        raise RuntimeError(
            f"XLSX exceeds compressed byte budget "
            f"({MAX_TABLE_BYTES // (1024 * 1024)} MiB); extraction stopped"
        )
    with zipfile.ZipFile(path) as zf:
        table_bytes = sum(
            info.file_size for info in zf.infolist()
            if info.filename == "xl/sharedStrings.xml"
            or re.fullmatch(r"xl/worksheets/sheet\d+\.xml", info.filename)
        )
    if table_bytes > MAX_TABLE_BYTES:
        raise RuntimeError(
            f"XLSX table XML exceeds uncompressed byte budget "
            f"({MAX_TABLE_BYTES // (1024 * 1024)} MiB); extraction stopped"
        )
    try:
        import openpyxl  # noqa: PLC0415

        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        parts = []
        budget = {
            "rows": MAX_TABLE_ROWS,
            "cells": MAX_TABLE_CELLS,
            "text": MAX_TABLE_TEXT_CHARS,
        }
        for ws in wb.worksheets:
            body = _collect_table(
                ws.iter_rows(values_only=True),
                max_rows,
                summary=True,
                budget=budget,
            )
            parts.append(f"## Sheet: {ws.title}\n\n" + body)
        return "\n\n".join(parts)
    except ImportError:
        return _xlsx_zip_fallback(path, max_rows)


def extract_csv(path: Path, max_rows: int) -> str:
    if path.stat().st_size > MAX_TABLE_BYTES:
        raise RuntimeError(
            f"table exceeds input byte budget "
            f"({MAX_TABLE_BYTES // (1024 * 1024)} MiB); extraction stopped"
        )
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    with open(path, newline="", encoding="utf-8", errors="replace") as fh:
        return _collect_table(csv.reader(fh, delimiter=delimiter), max_rows, summary=True)


def extract_docx(path: Path, max_rows: int, max_chars: int = DEFAULT_MAX_CHARS) -> str:
    try:
        with zipfile.ZipFile(path) as zf:
            _validate_ooxml_archive_bounded(zf, max_chars)
            _read_ooxml_member_bounded(zf, "word/document.xml", max_chars)
    except zipfile.BadZipFile:
        # Let python-docx or textutil report malformed legacy/test inputs.
        pass
    try:
        import docx  # noqa: PLC0415

        d = docx.Document(str(path))

        def chunks():
            for paragraph in d.paragraphs:
                if paragraph.text.strip():
                    yield paragraph.text
            for table in d.tables:
                rows = ([c.text for c in row.cells] for row in table.rows)
                yield _collect_table(rows, max_rows)

        return _join_bounded(chunks(), max_chars, "DOCX", "\n\n")
    except ImportError:
        pass
    if shutil.which("textutil"):
        returncode, stdout, _ = _run_text_command_bounded(
            ["textutil", "-convert", "txt", "-stdout", str(path)], max_chars, "DOCX"
        )
        if returncode == 0:
            return stdout
    with zipfile.ZipFile(path) as zf:
        return _xml_text(_read_ooxml_member_bounded(zf, "word/document.xml", max_chars))


def extract_pptx(path: Path, max_chars: int = DEFAULT_MAX_CHARS) -> str:
    with zipfile.ZipFile(path) as zf:
        slides = sorted(n for n in zf.namelist() if re.fullmatch(r"ppt/slides/slide\d+\.xml", n))
        return _join_bounded(
            (f"## Slide {i + 1}\n\n{_xml_text(_read_ooxml_member_bounded(zf, n, max_chars))}"
             for i, n in enumerate(slides)),
            max_chars,
            "PPTX",
            "\n\n",
        )


def extract_zip(path: Path, max_rows: int, member_cap: int = 20,
                total_budget: int = 64 * 1024 * 1024,
                _budget: dict[str, int] | None = None, _depth: int = 0,
                *, max_chars: int = DEFAULT_MAX_CHARS) -> str:
    # Archive → manifest + recursive extraction of the first N supported members.
    # RESOURCE BOUNDS (attachments are untrusted): cap the member count AND the
    # cumulative uncompressed bytes we materialize, so a zip-bomb / oversized
    # archive can't exhaust the host's disk or memory. A member whose declared
    # uncompressed size would blow the budget is skipped, not written.
    if _depth >= MAX_ARCHIVE_DEPTH:
        return f"[skipped — archive nesting exceeds {MAX_ARCHIVE_DEPTH} levels]"
    budget = _budget if _budget is not None else {
        "members": member_cap,
        "bytes": total_budget,
    }
    with zipfile.ZipFile(path) as zf:
        infos = [inf for inf in zf.infolist() if not inf.filename.endswith("/")]
        names = [inf.filename for inf in infos]
        parts = ["## Archive contents\n\n" + "\n".join(f"- {n}" for n in names[:200])]
        import tempfile  # noqa: PLC0415

        stopped = False
        processed = 0
        with tempfile.TemporaryDirectory() as td:
            for inf in infos:
                name = inf.filename
                if budget["members"] <= 0:
                    if _depth == 0 and processed == member_cap:
                        parts.append(
                            f"[doc-ingest: extracted first {member_cap} of "
                            f"{len(names)} members — shared archive member budget reached]"
                        )
                    else:
                        parts.append(f"[doc-ingest: shared archive member budget "
                                     f"({member_cap}) reached]")
                    stopped = True
                    break
                if inf.file_size > budget["bytes"]:
                    parts.append(f"## {name}\n\n[skipped — shared archive extraction budget "
                                 f"({total_budget // (1024 * 1024)} MiB) reached]")
                    stopped = True
                    break
                budget["members"] -= 1
                budget["bytes"] -= inf.file_size
                processed += 1
                target = Path(td) / Path(name).name  # flatten: zip-slip-safe, basename only
                target.write_bytes(zf.read(name))
                try:
                    kind, text = extract(
                        target, max_rows, max_chars,
                        _archive_budget=budget,
                        _archive_depth=_depth + 1,
                    )
                    if kind in {"image", "audio"}:
                        parts.append(f"## {name}\n\n[{kind} member — handled natively, not extracted]")
                    else:
                        parts.append(f"## {name} ({kind})\n\n{text}")
                except Exception as exc:  # noqa: BLE001 — one bad member must not sink the archive
                    parts.append(f"## {name}\n\n[extraction failed: {exc}]")
        if len(names) > member_cap and not stopped:
            parts.append(f"[doc-ingest: extracted first {member_cap} of {len(names)} members]")
    return "\n\n".join(parts)


def extract_textutil(path: Path, max_chars: int = DEFAULT_MAX_CHARS) -> str:
    if not shutil.which("textutil"):
        raise RuntimeError(f"no extractor for {path.suffix} (textutil unavailable on this host)")
    returncode, stdout, stderr = _run_text_command_bounded(
        ["textutil", "-convert", "txt", "-stdout", str(path)], max_chars, "textutil"
    )
    if returncode != 0:
        raise RuntimeError(f"textutil failed on {path.name}: {stderr.strip()[:200]}")
    return stdout


TABULAR_SUFFIXES = {".xlsx", ".xlsm", ".csv", ".tsv"}


def _stream_rows_to_csv(rows_iter, budget: Optional[dict[str, int]]) -> str:
    """Render rows to CSV incrementally under a shared fail-closed budget.

    `budget` is {"cells": N, "bytes": N}, mutated cumulatively so multi-sheet
    workbooks share ONE budget; None means the caller passed the explicit
    --csv-no-budget override. Exactness contract: on budget exhaustion this
    raises loudly — it never truncates, because a partial table computes a
    wrong aggregate.
    """
    parts: list[str] = []
    row_buf = io.StringIO()
    writer = csv.writer(row_buf)
    for raw in rows_iter:
        row = ["" if c is None else c for c in raw]
        row_buf.seek(0)
        row_buf.truncate()
        writer.writerow(row)
        chunk = row_buf.getvalue()
        if budget is not None:
            budget["cells"] -= len(row)
            # Budget the ENCODED size of what will be emitted, not
            # StringIO.tell() character deltas — multibyte and replacement
            # characters render more UTF-8 bytes than characters, and the
            # byte budget's contract is bytes-on-stdout (qingyun CR
            # 2026-07-31: a 31-char row encoding to 91 bytes must count 91).
            budget["bytes"] -= len(chunk.encode("utf-8"))
            if budget["cells"] < 0:
                raise RuntimeError(
                    f"--csv table exceeds the {MAX_TABLE_CELLS:,}-cell compute "
                    "budget (fail-closed: attachments are untrusted and "
                    "compute-exact mode never truncates); rerun with "
                    "--csv-no-budget to override")
            if budget["bytes"] < 0:
                raise RuntimeError(
                    f"--csv output exceeds the "
                    f"{CSV_MAX_OUTPUT_BYTES // (1024 * 1024)} MiB compute "
                    "budget (fail-closed: attachments are untrusted and "
                    "compute-exact mode never truncates); rerun with "
                    "--csv-no-budget to override")
        parts.append(chunk)
    return "".join(parts).rstrip("\n")


def extract_table_csv(path: Path, unbounded: bool = False) -> str:
    """Exact, uncapped per-sheet CSV for COMPUTING over — the machine twin of the
    row-capped markdown view (which is for reading/summaries).

    Quantitative answers must come from computation over the exact table, not
    from eyeballing extracted text — validated on the GAIA file-attached subset
    (3/3 multi-cell numeric misses flipped, 84.2%→92.1%, 2026-07-30). Two
    exactness rules follow from that purpose:
      * no row caps and no char caps by default — a silently truncated table
        computes a silently wrong aggregate;
      * xlsx without openpyxl is REFUSED (clear error) rather than served by
        the approximate zip-XML fallback — approximate cells are fine to read,
        not to compute with;
      * no caps does NOT mean no bounds: attachments are untrusted, so inputs
        over CSV_MAX_SOURCE_BYTES and renders over CSV_MAX_OUTPUT_BYTES /
        MAX_TABLE_CELLS fail LOUDLY (never truncate) unless the caller passes
        the explicit --csv-no-budget override.
    """
    suffix = path.suffix.lower()
    if suffix not in TABULAR_SUFFIXES:
        raise ValueError(f"--csv applies to tabular files ({', '.join(sorted(TABULAR_SUFFIXES))}); got {suffix or 'no suffix'}")
    if not unbounded:
        size = path.stat().st_size
        if size > CSV_MAX_SOURCE_BYTES:
            raise RuntimeError(
                f"--csv input is {size:,} bytes, over the "
                f"{CSV_MAX_SOURCE_BYTES // (1024 * 1024)} MiB compute budget "
                "(fail-closed: attachments are untrusted and compute-exact "
                "mode never truncates); rerun with --csv-no-budget to override")
    budget = None if unbounded else {"cells": MAX_TABLE_CELLS,
                                     "bytes": CSV_MAX_OUTPUT_BYTES}
    if suffix in {".csv", ".tsv"}:
        delimiter = "\t" if suffix == ".tsv" else ","
        with open(path, newline="", encoding="utf-8", errors="replace") as fh:
            return _stream_rows_to_csv(csv.reader(fh, delimiter=delimiter), budget)
    # XLSX/XLSM path. The compressed-size gate above cannot catch an OOXML zip
    # bomb (a small archive whose sharedStrings/worksheet XML inflates far past
    # the budget) — openpyxl.load_workbook would allocate past the compute budget
    # before any cell/byte budget could fire. Preflight the UNCOMPRESSED archive
    # size here, BEFORE importing/handing the file to openpyxl, so the bomb is
    # rejected even where openpyxl is absent. Mirrors extract_xlsx's
    # _validate_ooxml_archive_bounded (qingyun CR #2434). --csv-no-budget opts out.
    if not unbounded:
        try:
            with zipfile.ZipFile(path) as zf:
                uncompressed = sum(info.file_size for info in zf.infolist())
        except zipfile.BadZipFile:
            uncompressed = None  # not a zip — let openpyxl raise its own clear error
        if uncompressed is not None and uncompressed > CSV_MAX_SOURCE_BYTES:
            raise RuntimeError(
                f"--csv xlsx inflates to {uncompressed:,} uncompressed bytes, over "
                f"the {CSV_MAX_SOURCE_BYTES // (1024 * 1024)} MiB compute budget "
                "(fail-closed: attachments are untrusted and compute-exact mode "
                "never truncates); rerun with --csv-no-budget to override")
    try:
        import openpyxl  # noqa: PLC0415
    except ImportError as exc:
        raise RuntimeError("--csv for xlsx needs openpyxl (the zip-XML fallback is "
                           "approximate — fine for reading, not for computing)") from exc
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"=== sheet: {ws.title} ===\n"
                     + _stream_rows_to_csv(ws.iter_rows(values_only=True), budget))
    return "\n\n".join(parts)


def extract(path: Path, max_rows: int, max_chars: int = DEFAULT_MAX_CHARS,
            _archive_budget: dict[str, int] | None = None,
            _archive_depth: int = 0) -> tuple[str, str]:
    """Returns (kind, text). Raises on failure; special kinds 'image'/'audio' carry a pointer."""
    suffix = path.suffix.lower()
    if suffix in IMAGE_SUFFIXES:
        return "image", "doc-ingest: images are read natively by the agent (Read tool) — no extraction here."
    if suffix in AUDIO_SUFFIXES:
        return "audio", "doc-ingest: use skills/audio-transcribe for audio files."
    if suffix == ".pdf":
        return "pdf", extract_pdf(path, max_chars)
    if suffix in {".xlsx", ".xlsm"}:
        return "xlsx", extract_xlsx(path, max_rows)
    if suffix in {".csv", ".tsv"}:
        return "table", extract_csv(path, max_rows)
    if suffix == ".docx":
        return "docx", extract_docx(path, max_rows, max_chars)
    if suffix == ".pptx":
        return "pptx", extract_pptx(path, max_chars)
    if suffix == ".zip":
        return "zip", extract_zip(
            path, max_rows,
            _budget=_archive_budget,
            _depth=_archive_depth,
            max_chars=max_chars,
        )
    if suffix in {".rtf", ".doc"}:
        return "textutil", extract_textutil(path, max_chars)
    if suffix in TEXT_SUFFIXES or not suffix:
        return "text", _read_text_bounded(path, max_chars)
    # Unknown suffix: try text read — better a replaced-chars dump than a refusal.
    return "text?", _read_text_bounded(path, max_chars)


def main(argv: list[str]) -> int:
    args = list(argv)
    as_json = "--json" in args
    if as_json:
        args.remove("--json")
    csv_mode = "--csv" in args
    if csv_mode:
        args.remove("--csv")
    csv_unbounded = "--csv-no-budget" in args
    if csv_unbounded:
        args.remove("--csv-no-budget")
    max_chars, max_rows = DEFAULT_MAX_CHARS, DEFAULT_MAX_ROWS
    explicit_chars = False
    for flag, setter in (("--max-chars", "chars"), ("--max-rows", "rows")):
        if flag in args:
            i = args.index(flag)
            try:
                value = int(args[i + 1])
            except (IndexError, ValueError):
                print(f"doc-ingest: {flag} needs an integer", file=sys.stderr)
                return 2
            if setter == "chars":
                max_chars = value
                explicit_chars = True
            else:
                max_rows = value
            del args[i:i + 2]
    # --csv is the compute-exact view: the default char cap must not silently
    # truncate a table mid-row (a truncated table computes a wrong aggregate).
    # An EXPLICIT --max-chars still wins — the caller asked for it.
    if csv_mode and not explicit_chars:
        max_chars = 0
    if not args:
        print("usage: ingest.py <file> [<file> ...] [--json] [--csv] [--csv-no-budget] [--max-chars N] [--max-rows N]", file=sys.stderr)
        return 2

    had_failure = had_pointer = False
    for name in args:
        path = Path(name)
        if not path.is_file():
            had_failure = True
            result = {"file": name, "kind": "missing", "ok": False, "error": "not a file"}
        else:
            try:
                if csv_mode:
                    kind, text = "csv", extract_table_csv(path, unbounded=csv_unbounded)
                else:
                    kind, text = extract(path, max_rows, max_chars)
                if kind in {"image", "audio"}:
                    had_pointer = True
                    result = {"file": name, "kind": kind, "ok": False, "error": text}
                else:
                    result = {"file": name, "kind": kind, "ok": True, "text": _truncate(text, max_chars)}
            except Exception as exc:  # noqa: BLE001 — every per-file failure must be reported, not raised
                had_failure = True
                result = {"file": name, "kind": "error", "ok": False, "error": str(exc)}
        if as_json:
            print(json.dumps(result, ensure_ascii=False))
        elif result["ok"]:
            header = f"===== {name} ({result['kind']}) =====\n" if len(args) > 1 else ""
            print(header + result["text"])
        else:
            print(f"doc-ingest: {name}: {result['error']}", file=sys.stderr)

    # Aggregate exit status must reflect the per-file results: any hard failure
    # dominates (1); otherwise any pointer-only result (image/audio, ok:false)
    # yields 3 — even in a mixed batch, so a caller never reads exit 0 while a
    # record is ok:false. All-extracted → 0.
    if had_failure:
        return 1
    return 3 if had_pointer else 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))

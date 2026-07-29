#!/usr/bin/env python3
"""doc-ingest extraction tests (skills/doc-ingest/scripts/ingest.py).

Hermetic: every fixture is generated into a tempdir at run time; formats whose
optional library is missing on this host are exercised through their
dependency-free fallback (pptx/xlsx zip-XML paths) or skipped with a notice.
Run: python3 tests/doc-ingest.test.py
"""
import importlib.util
import io
import json
import sys
import tempfile
import types
import zipfile
from contextlib import redirect_stderr, redirect_stdout
from pathlib import Path
from unittest import mock

_ROOT = Path(__file__).resolve().parent.parent
spec = importlib.util.spec_from_file_location(
    "ingest", _ROOT / "skills" / "doc-ingest" / "scripts" / "ingest.py")
ingest = importlib.util.module_from_spec(spec)
spec.loader.exec_module(ingest)

passed = []


def run_cli(argv):
    out, err = io.StringIO(), io.StringIO()
    with redirect_stdout(out), redirect_stderr(err):
        code = ingest.main(argv)
    return code, out.getvalue(), err.getvalue()


def check(name, condition, detail=""):
    assert condition, f"FAIL {name}: {detail}"
    passed.append(name)


def make_pptx(path: Path):
    # Minimal OOXML shell — just enough structure for the slide-XML scrape.
    slide = (b'<?xml version="1.0"?><p:sld xmlns:p="x" xmlns:a="y">'
             b"<p:txBody><a:p><a:r><a:t>Hello from slide one</a:t></a:r></a:p></p:txBody></p:sld>")
    with zipfile.ZipFile(path, "w") as zf:
        zf.writestr("ppt/slides/slide1.xml", slide)


with tempfile.TemporaryDirectory() as td:
    tmp = Path(td)

    # 1. Plain text passes through, and --max-chars truncates with a notice.
    txt = tmp / "note.txt"
    txt.write_text("alpha beta gamma " * 100)
    code, out, _ = run_cli([str(txt)])
    check("text-read", code == 0 and "alpha beta" in out)
    code, out, _ = run_cli([str(txt), "--max-chars", "50"])
    check("truncation-notice", "truncated at 50 chars" in out, out[-120:])

    # 2. CSV → markdown table, row cap appends a visible notice.
    csvf = tmp / "data.csv"
    csvf.write_text("name,qty\nwidget,3\ngadget,5\nsprocket,7\n")
    code, out, _ = run_cli([str(csvf)])
    check("csv-table", code == 0 and "| name | qty |" in out and "| gadget | 5 |" in out)
    code, out, _ = run_cli([str(csvf), "--max-rows", "2"])
    check("csv-row-cap", "showing 2 of 4 rows" in out, out[-120:])

    # 3. PPTX via the dependency-free zip-XML path (always runs).
    pptx = tmp / "deck.pptx"
    make_pptx(pptx)
    code, out, _ = run_cli([str(pptx)])
    check("pptx-slide-text", code == 0 and "Hello from slide one" in out and "Slide 1" in out)

    # 4. XLSX when openpyxl is available (else skip — fallback covered by pptx path).
    try:
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Inventory"
        ws.append(["part", "count"])
        ws.append(["bolt", 42])
        xlsx = tmp / "inv.xlsx"
        wb.save(str(xlsx))
        code, out, _ = run_cli([str(xlsx)])
        check("xlsx-sheet", code == 0 and "Sheet: Inventory" in out and "| bolt | 42 |" in out)
    except ImportError:
        print("SKIP xlsx-sheet (openpyxl not installed)")

    # 5. DOCX when python-docx is available (else skip) — with a table too.
    try:
        import docx

        d = docx.Document()
        d.add_paragraph("Quarterly summary paragraph.")
        table = d.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "region"
        table.rows[0].cells[1].text = "sales"
        table.rows[1].cells[0].text = "west"
        table.rows[1].cells[1].text = "10"
        docxf = tmp / "report.docx"
        d.save(str(docxf))
        code, out, _ = run_cli([str(docxf)])
        check("docx-paragraph",
              code == 0 and "Quarterly summary paragraph." in out
              and "| region | sales |" in out)
    except ImportError:
        print("SKIP docx-paragraph (python-docx not installed)")

    # 6. Image → pointer to native reading, exit 3, nothing on stdout.
    img = tmp / "shot.png"
    img.write_bytes(b"\x89PNG\r\n\x1a\n")
    code, out, err = run_cli([str(img)])
    check("image-pointer", code == 3 and not out.strip() and "Read tool" in err, f"code={code}")

    # 7. Audio → pointer to audio-transcribe.
    audio = tmp / "memo.mp3"
    audio.write_bytes(b"ID3")
    code, out, err = run_cli([str(audio)])
    check("audio-pointer", code == 3 and "audio-transcribe" in err, f"code={code}")

    # 8. Missing file → exit 1 with a stderr report.
    code, out, err = run_cli([str(tmp / "ghost.pdf")])
    check("missing-file", code == 1 and "not a file" in err, f"code={code}")

    # 9. Corrupt PDF → reported as a per-file error, never a traceback.
    bad = tmp / "bad.pdf"
    bad.write_text("this is not a pdf")
    code, out, err = run_cli([str(bad)])
    check("corrupt-pdf-reported", code == 1 and "bad.pdf" in err, f"code={code} err={err[:120]}")

    # 10. --json mode emits one machine-readable line per input, mixed statuses.
    code, out, err = run_cli([str(txt), str(img), "--json"])
    lines = [json.loads(line) for line in out.strip().splitlines()]
    check("json-lines",
          len(lines) == 2 and lines[0]["ok"] is True and lines[1]["ok"] is False
          and lines[1]["kind"] == "image",
          out[:200])

    # 11. Bad invocation (no files) → usage + exit 2.
    code, out, err = run_cli([])
    check("usage-exit-2", code == 2 and "usage:" in err)

    # 12. ZIP → manifest + recursive member extraction; image members deferred.
    zipf = tmp / "bundle.zip"
    with zipfile.ZipFile(zipf, "w") as zf:
        zf.writestr("inner/data.csv", "a,b\n1,2\n")
        zf.writestr("readme.txt", "hello archive")
        zf.writestr("pic.png", "\x89PNG")
    code, out, _ = run_cli([str(zipf)])
    check("zip-recursive",
          code == 0 and "Archive contents" in out and "| a | b |" in out
          and "hello archive" in out and "handled natively" in out,
          out[:300])

    # 13. Empty table → sentinel (helper branch).
    check("empty-table", ingest._rows_to_markdown([], 500) == "(empty table)")

    # 14. Unknown suffix → best-effort text read ("text?" kind).
    weird = tmp / "note.xyz"
    weird.write_text("payload body")
    kind, text = ingest.extract(weird, 500)
    check("unknown-suffix", kind == "text?" and "payload body" in text)

    # 15. PDF with no extractor available → RuntimeError naming the gap.
    #     Force pdftotext absent; pypdf/fitz import is stubbed to fail.
    pdf = tmp / "x.pdf"
    pdf.write_bytes(b"%PDF-1.4 not-really")
    with mock.patch.object(ingest.shutil, "which", return_value=None), \
         mock.patch.dict(sys.modules, {"pypdf": None, "fitz": None}):
        try:
            ingest.extract_pdf(pdf)
            check("pdf-no-extractor", False, "expected RuntimeError")
        except RuntimeError as exc:
            check("pdf-no-extractor", "no PDF extractor" in str(exc))

    # 16. PDF pypdf fallback path: inject a fake pypdf whose reader yields text.
    fake_pypdf = types.ModuleType("pypdf")
    fake_pypdf.PdfReader = lambda _p: types.SimpleNamespace(
        pages=[types.SimpleNamespace(extract_text=lambda: "page one text")])
    with mock.patch.object(ingest.shutil, "which", return_value=None), \
         mock.patch.dict(sys.modules, {"pypdf": fake_pypdf}):
        check("pdf-pypdf-fallback", "page one text" in ingest.extract_pdf(pdf))

    # 17. XLSX fallback (openpyxl missing) → shared-strings zip scrape.
    fake_xlsx = tmp / "f.xlsx"
    with zipfile.ZipFile(fake_xlsx, "w") as zf:
        zf.writestr("xl/sharedStrings.xml",
                    '<sst><si><t>CellFromSharedStrings</t></si></sst>')
    with mock.patch.dict(sys.modules, {"openpyxl": None}):
        check("xlsx-fallback",
              "CellFromSharedStrings" in ingest.extract_xlsx(fake_xlsx, 500))

    # 18. DOCX fallback (python-docx + textutil missing) → document.xml scrape.
    fake_docx = tmp / "f.docx"
    with zipfile.ZipFile(fake_docx, "w") as zf:
        zf.writestr("word/document.xml",
                    '<w:document><w:body><w:p >ParaFromXml</w:p></w:body></w:document>')
    with mock.patch.dict(sys.modules, {"docx": None}), \
         mock.patch.object(ingest.shutil, "which", return_value=None):
        check("docx-fallback", "ParaFromXml" in ingest.extract_docx(fake_docx, 500))

    # 19. RTF via textutil unavailable → clear RuntimeError.
    rtf = tmp / "x.rtf"
    rtf.write_text("{\\rtf1 hi}")
    with mock.patch.object(ingest.shutil, "which", return_value=None):
        try:
            ingest.extract_textutil(rtf)
            check("textutil-missing", False, "expected RuntimeError")
        except RuntimeError as exc:
            check("textutil-missing", "textutil unavailable" in str(exc))

    # 20. --max-rows flag parsed; bad value → exit 2.
    manyrows = tmp / "big.csv"
    manyrows.write_text("h\n" + "\n".join(str(i) for i in range(10)) + "\n")
    code, out, _ = run_cli([str(manyrows), "--max-rows", "3"])
    check("max-rows-flag", code == 0 and "showing 3 of" in out, out[-120:])
    code, _, err = run_cli([str(manyrows), "--max-rows", "notanint"])
    check("max-rows-bad", code == 2 and "needs an integer" in err)

    # 21. ZIP member cap: >cap members truncates with a notice.
    bigzip = tmp / "many.zip"
    with zipfile.ZipFile(bigzip, "w") as zf:
        for i in range(25):
            zf.writestr(f"f{i}.txt", f"content {i}")
    text = ingest.extract_zip(bigzip, 500, member_cap=20)
    check("zip-member-cap", "extracted first 20 of 25 members" in text, text[-160:])

    # 22. ZIP member whose extraction raises → reported inline, archive survives.
    badzip = tmp / "badmember.zip"
    with zipfile.ZipFile(badzip, "w") as zf:
        zf.writestr("ok.txt", "fine")
        zf.writestr("broken.pdf", "not a real pdf")  # forces extract_pdf failure
    with mock.patch.object(ingest.shutil, "which", return_value=None), \
         mock.patch.dict(sys.modules, {"pypdf": None, "fitz": None}):
        text = ingest.extract_zip(badzip, 500)
    check("zip-bad-member", "fine" in text and "extraction failed" in text, text[-200:])

    # 23. RTF/.doc dispatch → textutil success path (mock which + subprocess).
    rtf2 = tmp / "y.rtf"
    rtf2.write_text("{\\rtf1 body}")
    fake_proc = types.SimpleNamespace(returncode=0, stdout="converted rtf text", stderr="")
    with mock.patch.object(ingest.shutil, "which", return_value="/usr/bin/textutil"), \
         mock.patch.object(ingest.subprocess, "run", return_value=fake_proc):
        kind, text = ingest.extract(rtf2, 500)
    check("rtf-textutil-success", kind == "textutil" and "converted rtf text" in text)

    # 24. textutil non-zero return → RuntimeError naming the failure.
    fail_proc = types.SimpleNamespace(returncode=1, stdout="", stderr="boom")
    with mock.patch.object(ingest.shutil, "which", return_value="/usr/bin/textutil"), \
         mock.patch.object(ingest.subprocess, "run", return_value=fail_proc):
        try:
            ingest.extract_textutil(rtf2)
            check("textutil-nonzero", False, "expected RuntimeError")
        except RuntimeError as exc:
            check("textutil-nonzero", "textutil failed" in str(exc))

print(f"OK — {len(passed)} checks passed: {', '.join(passed)}")
